"""
广州大学毕业论文（设计）格式修正工具

基于学校模板自动修正论文格式：
- 页面设置（封面/正文页边距）
- 字体字号（宋体/黑体/Times New Roman，三号/四号/小四/五号）
- 段落格式（首行缩进2字符、行距固定23pt）
- 标题样式（章标题、节标题、小节标题）
- 摘要/关键词格式
- 英文摘要格式
- 参考文献格式（悬挂缩进）
- 图表题注格式
- 页眉页脚格式
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ── 单位转换常量 ──
# 1 inch = 914400 EMU
# 1 cm = 360000 EMU
# 1 pt = 12700 EMU

# ── 模板格式常量 ──
COVER_MARGINS = {"top": Cm(2.5), "bottom": Cm(2.5), "left": Cm(3.0), "right": Cm(2.6)}
BODY_MARGINS = {"top": Cm(2.5), "bottom": Cm(2.5), "left": Cm(3.2), "right": Cm(3.2)}

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_TIMES = "Times New Roman"
FONT_KAI = "楷体_GB2312"

SIZE_SANHAO = Pt(16)   # 三号
SIZE_SIHAO = Pt(14)    # 四号
SIZE_XIAOSI = Pt(12)   # 小四
SIZE_WUHAO = Pt(10.5)  # 五号
SIZE_XIAOWU = Pt(9)    # 小五

BODY_LINE_SPACING = 1.5   # 1.5倍行距
BODY_INDENT = Pt(24)      # 首行缩进2字符（约24pt）


def set_run_font(run, cn_font=None, en_font=None, size=None, bold=None, color=None):
    """设置 run 的中英文字体、字号、加粗、颜色。"""
    if cn_font:
        run.font.name = cn_font
        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
            r.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), cn_font)
        if en_font:
            rFonts.set(qn('w:ascii'), en_font)
            rFonts.set(qn('w:hAnsi'), en_font)
    if en_font and not cn_font:
        run.font.name = en_font
    if size:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_paragraph_format(paragraph, alignment=None, line_spacing=None,
                          line_spacing_rule=None,
                          first_line_indent=None, space_before=None, space_after=None,
                          keep_with_next=None, outline_level=None):
    """设置段落格式。"""
    pf = paragraph.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if line_spacing is not None:
        pf.line_spacing = line_spacing
        # 整数行距（twips）→固定值；浮点行距（倍数）→多倍行距
        if line_spacing_rule is None and isinstance(line_spacing, (int, float)):
            if isinstance(line_spacing, int):
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            else:
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if keep_with_next is not None:
        pf.keep_with_next = keep_with_next
    if outline_level is not None:
        pf.outline_level = outline_level


def is_heading_paragraph(text):
    """判断段落是否为标题。"""
    # 章标题：纯中文如"前言"、"第一章"、"第1章"、"总结与展望"、"1 前言"、"3 系统需求分析"
    chapter_patterns = [
        r'^前言$', r'^绪论$', r'^结论$', r'^致谢$', r'^参考文献$',
        r'^第[一二三四五六七八九十\d]+章',
        r'^[一二三四五六七八九十]+、',  # "一、研究背景"
        r'^总结与展望$', r'^结束语$',
        r'^[1-9]\d*\s+.+',  # "1 前言", "2 相关理论基础"
    ]
    for pat in chapter_patterns:
        if re.match(pat, text):
            return "chapter"
    return None


def remove_empty_paragraphs(doc):
    """移除空白段落。"""
    paragraphs_to_remove = []
    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            paragraphs_to_remove.append(p)
    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)


def format_page_setup(doc):
    """修正页面设置（保留横向页面不做修改）。"""
    sections = doc.sections
    if len(sections) >= 1:
        sec = sections[0]
        # 只修改竖向页面，跳过横向页面
        if sec.page_width < sec.page_height:
            sec.page_width = Cm(21.0)
            sec.page_height = Cm(29.7)
            sec.top_margin = COVER_MARGINS["top"]
            sec.bottom_margin = COVER_MARGINS["bottom"]
            sec.left_margin = COVER_MARGINS["left"]
            sec.right_margin = COVER_MARGINS["right"]
    if len(sections) >= 2:
        sec = sections[1]
        if sec.page_width < sec.page_height:
            sec.page_width = Cm(21.0)
            sec.page_height = Cm(29.7)
            sec.top_margin = BODY_MARGINS["top"]
            sec.bottom_margin = BODY_MARGINS["bottom"]
            sec.left_margin = BODY_MARGINS["left"]
            sec.right_margin = BODY_MARGINS["right"]


def get_structure_zones(doc):
    """分析文档结构，返回各区域的段落索引范围。"""
    paragraphs = doc.paragraphs
    zones = {
        "cover": [],           # 封面
        "abstract_cn": [],     # 中文摘要
        "abstract_en": [],     # 英文摘要
        "toc": [],             # 目录
        "body": [],            # 正文
        "references": [],      # 参考文献
        "acknowledgement": [], # 致谢
    }

    current_zone = "cover"
    found_abstract = False
    found_abstract_en = False
    found_toc = False
    found_first_chapter = False
    found_references = False
    found_ack = False

    for i, p in enumerate(paragraphs):
        text = p.text.strip()

        # 检测中文摘要
        if text == "摘要" or (text.startswith("摘要") and not found_abstract):
            current_zone = "abstract_cn"
            found_abstract = True
        # 检测英文摘要
        elif re.match(r'^ABSTRACT', text, re.IGNORECASE) and not found_abstract_en:
            current_zone = "abstract_en"
            found_abstract_en = True
        # 检测目录
        elif text == "目  录" or text == "目录" and not found_toc:
            current_zone = "toc"
            found_toc = True
        # 检测正文开始（第一个章标题）
        elif not found_first_chapter and is_heading_paragraph(text) == "chapter":
            current_zone = "body"
            found_first_chapter = True
        # 检测参考文献
        elif text == "参考文献" and current_zone == "body":
            current_zone = "references"
            found_references = True
        # 检测致谢
        elif text == "致谢":
            current_zone = "acknowledgement"
            found_ack = True

        zones.setdefault(current_zone, []).append(i)

    return zones


SIZE_YIHAO = Pt(26)    # 一号

def format_cover(doc, zones):
    """修正封面格式（含表格）。"""
    cover_indices = zones.get("cover", [])

    for idx in cover_indices:
        p = doc.paragraphs[idx]
        text = p.text.strip()
        if not text:
            continue

        # 论文标题（"本 科 毕 业 论 文（设计）"）→ 楷体一号居中
        if ("毕业" in text.replace(" ", "") and "论文" in text.replace(" ", "") and "设计" in text.replace(" ", "")):
            for run in p.runs:
                set_run_font(run, cn_font=FONT_KAI, en_font=FONT_TIMES,
                            size=SIZE_YIHAO, bold=None)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            continue

        # "教 务 处 制" → 楷体三号居中
        if re.match(r'^教\s*务\s*处\s*制$', text):
            for run in p.runs:
                set_run_font(run, cn_font=FONT_KAI, en_font=FONT_TIMES,
                            size=SIZE_SANHAO)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            continue

        # 论文具体标题 → 黑体三号居中
        if len(text) >= 8 and not any(x in text for x in ["务", "处", "制", "班", "指导"]):
            for run in p.runs:
                set_run_font(run, cn_font=FONT_HEI, en_font=FONT_TIMES,
                            size=SIZE_SANHAO)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            continue

    # 修正封面表格（第一张表格）
    if doc.tables:
        cover_table = doc.tables[0]
        for row in cover_table.rows:
            cells = row.cells
            if len(cells) >= 2:
                # 左列：标签列 → 楷体小二号(18pt) 右对齐，不加粗
                for p in cells[0].paragraphs:
                    for run in p.runs:
                        set_run_font(run, cn_font=FONT_KAI, en_font=FONT_TIMES,
                                    size=Pt(18), bold=False)
                    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                                        line_spacing=BODY_LINE_SPACING)
                # 右列：内容列 → 楷体三号(16pt) 居中，不加粗
                for p in cells[1].paragraphs:
                    for run in p.runs:
                        set_run_font(run, cn_font=FONT_KAI, en_font=FONT_TIMES,
                                    size=SIZE_SANHAO, bold=False)
                    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                        line_spacing=BODY_LINE_SPACING)


def _is_chapter_style(style_name, text):
    """判断是否为章标题样式。"""
    if style_name in ("章标题", "一级标题"):
        return True
    if is_heading_paragraph(text) == "chapter":
        return True
    return False


def _is_numbered_heading(text, pattern):
    """检测 Normal 段落是否为编号标题（如 2.1 工程概况 或 2.1.1 地形、地貌）。

    标题特征：数字编号 + 中文/英文标题文字。排除纯数值+单位（如 0.28 mm）。
    """
    m = re.match(pattern, text)
    if not m:
        return False

    # 编号后的内容
    after_num = text[m.end():].strip()
    if not after_num:
        return False

    # 纯数值+单位（0.28 mm, 1.5 米）→ 不是标题
    if re.match(r'^[\d.\-]+\s*(mm|cm|m|km|MPa|kPa|kN|kg|t|℃|%|°|m/s|米|吨)\s*$', after_num):
        return False
    if re.match(r'^\d+(\.\d+)?\s*(mm|cm|m|km|MPa|kPa|kN|kg|t|℃|%|°|m/s)\b', after_num):
        return False

    # 必须有中文内容（实体的标题必然是中文）
    if not re.search(r'[\u4e00-\u9fff]', after_num):
        return False

    return True


def _is_section_style(style_name, text=""):
    """判断是否为节标题（1.1 工程概况 或 1.1工程概况）。"""
    if style_name in ("一级条标题", "二级标题"):
        return True
    if style_name == "Normal" and _is_numbered_heading(text, r'^\d+\.\d+\s*'):
        return True
    return False


def _is_subsection_style(style_name, text=""):
    """判断是否为小节标题（1.1.1 地形、地貌 或 1.1.1地形、地貌）。"""
    if style_name in ("二级条标题", "三级标题"):
        return True
    if style_name == "Normal" and _is_numbered_heading(text, r'^\d+\.\d+\.\d+\s*'):
        return True
    return False


def _is_caption_style(style_name):
    """判断是否为图表题注样式。"""
    return style_name in ("图和表", "Caption", "Normal (Web)")


def format_body_paragraph(p, is_reference=False):
    """修正正文段落格式。"""
    text = p.text.strip()
    if not text:
        return

    style_name = p.style.name

    # ── 章标题（添加段前分页，每章新起一页） ──
    if _is_chapter_style(style_name, text):
        for run in p.runs:
            set_run_font(run, cn_font=FONT_HEI, en_font=FONT_TIMES,
                        size=SIZE_SANHAO, bold=None)
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            line_spacing=BODY_LINE_SPACING,
                            space_before=Pt(3), space_after=Pt(3))
        # 添加段前分页（"参考文献"和"致谢"除外，它们通常紧跟正文不强制分页）
        if text not in ("参考文献", "致谢"):
            pPr = p._element.get_or_add_pPr()
            pb = pPr.find(qn('w:pageBreakBefore'))
            if pb is None:
                pb = parse_xml(f'<w:pageBreakBefore {nsdecls("w")} />')
                pPr.append(pb)
        return

    # ── 节标题（如"1.1 XXX"） ──
    if _is_section_style(style_name, text):
        for run in p.runs:
            set_run_font(run, cn_font=FONT_HEI, en_font=FONT_TIMES, size=SIZE_XIAOSI)
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            line_spacing=BODY_LINE_SPACING,
                            first_line_indent=None)
        return

    # ── 小节标题（如"1.1.1 XXX"） ──
    if _is_subsection_style(style_name, text):
        for run in p.runs:
            set_run_font(run, cn_font=FONT_HEI, en_font=FONT_TIMES, size=SIZE_XIAOSI)
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            line_spacing=BODY_LINE_SPACING,
                            first_line_indent=None)
        return

    # ── 图和表（含内容模式匹配：以"图"或"表"开头+数字的段落） ──
    is_caption = _is_caption_style(style_name) or bool(
        re.match(r'^(图|表)\s*[\d.]+[\s\S]', text)
    )
    if is_caption:
        for run in p.runs:
            set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TIMES, size=SIZE_WUHAO)
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            line_spacing=BODY_LINE_SPACING,
                            first_line_indent=None)
        return

    # ── 公式 ──
    if style_name == "公式":
        for run in p.runs:
            set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TIMES, size=SIZE_XIAOSI)
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            line_spacing=BODY_LINE_SPACING)
        return

    # ── 参考文献条目 ──
    if is_reference or style_name == "参考文献":
        for run in p.runs:
            set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TIMES, size=SIZE_XIAOSI)
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            line_spacing=BODY_LINE_SPACING,
                            first_line_indent=None)
        return

    # ── 代码块（Consolas 等宽字体，不修改） ──
    has_mono = any(r.font.name and r.font.name in ("Consolas", "Courier New", "monospace")
                   for r in p.runs)
    if has_mono:
        return

    # ── 普通正文 ──
    for run in p.runs:
        if not run.font.name and run.font.size is None:
            set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TIMES, size=SIZE_XIAOSI)

    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        line_spacing=BODY_LINE_SPACING,
                        first_line_indent=BODY_INDENT)


def _add_blank_before_headings(doc, zones):
    """确保每个标题前都有一个空行（章/节/小节标题前均需空行）。"""
    body_indices = zones.get("body", [])
    if not body_indices:
        return

    paragraphs = list(doc.paragraphs)  # snapshot
    body_elem = paragraphs[body_indices[0]]._element.getparent()

    # 从后往前处理，避免插入后索引失效
    heading_positions = []
    for idx in body_indices:
        p = paragraphs[idx]
        text = p.text.strip()
        if not text:
            continue
        s = p.style.name

        is_heading = _is_chapter_style(s, text) or _is_section_style(s, text) \
                     or _is_subsection_style(s, text) or _is_caption_style(s)
        if not is_heading:
            continue

        if idx > 0:
            prev_p = paragraphs[idx - 1]
            if prev_p.text.strip():
                heading_positions.append(idx)

    # 从后往前插入，避免索引偏移
    for idx in reversed(heading_positions):
        p = paragraphs[idx]
        if p._element.getparent() is not None:
            empty_p = OxmlElement('w:p')
            body_elem.insert(
                list(body_elem).index(p._element),
                empty_p
            )


def format_body(doc, zones):
    """修正正文格式。"""
    body_indices = zones.get("body", [])
    ref_indices = zones.get("references", [])

    for idx in body_indices:
        p = doc.paragraphs[idx]
        format_body_paragraph(p)

    for idx in ref_indices:
        p = doc.paragraphs[idx]
        format_body_paragraph(p, is_reference=True)


def format_abstract(doc, zones):
    """修正摘要格式。"""
    cn_indices = zones.get("abstract_cn", [])
    en_indices = zones.get("abstract_en", [])

    for idx in cn_indices:
        p = doc.paragraphs[idx]
        text = p.text.strip()
        if not text:
            continue

        # "摘要"标题行（标题词黑体四号，正文内容宋体小四不加粗）
        if text == "摘要" or text.startswith("摘要"):
            for run in p.runs:
                run_text = run.text
                if re.match(r'^摘要\s*$', run_text) or run_text == "摘要":
                    set_run_font(run, cn_font=FONT_HEI, en_font=FONT_TIMES,
                                size=SIZE_SIHAO)
                else:
                    set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TIMES,
                                size=SIZE_XIAOSI, bold=False)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                line_spacing=BODY_LINE_SPACING)
            continue

        # "关键词"标题行（标题词黑体四号，关键词内容宋体小四）
        if text.startswith("关键词") or text.startswith("关键词"):
            for run in p.runs:
                run_text = run.text
                if re.match(r'^关键词\s*$', run_text) or run_text == "关键词":
                    set_run_font(run, cn_font=FONT_HEI, en_font=FONT_TIMES,
                                size=SIZE_SIHAO)
                else:
                    set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TIMES,
                                size=SIZE_XIAOSI, bold=False)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                line_spacing=BODY_LINE_SPACING)
            continue

        # 摘要正文内容
        for run in p.runs:
            set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TIMES, size=SIZE_XIAOSI)
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            line_spacing=BODY_LINE_SPACING,
                            first_line_indent=BODY_INDENT)

    for idx in en_indices:
        p = doc.paragraphs[idx]
        text = p.text.strip()
        if not text:
            continue

        # ABSTRACT / KEY WORDS 标题行（标题词加粗，内容不加粗）
        is_abst_title = re.match(r'^(ABSTRACT|KEY\s*WORDS)', text, re.IGNORECASE)
        if is_abst_title:
            abst_kw = is_abst_title.group(1).upper()
            for run in p.runs:
                run_text = run.text
                # 判断这个 run 是否只是标题词本身
                is_title_word = bool(re.match(r'^(ABSTRACT|KEY|WORDS|KEY\s*WORDS)\s*$',
                                              run_text, re.IGNORECASE))
                if is_title_word:
                    set_run_font(run, en_font=FONT_TIMES,
                                size=SIZE_SIHAO, bold=True)
                else:
                    set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TIMES,
                                size=SIZE_XIAOSI, bold=False)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                line_spacing=BODY_LINE_SPACING)
            continue

        # 英文摘要正文（不加粗）
        for run in p.runs:
            set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TIMES,
                        size=SIZE_XIAOSI, bold=False)
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            line_spacing=BODY_LINE_SPACING,
                            first_line_indent=BODY_INDENT)


def format_acknowledgement(doc, zones):
    """修正致谢格式。"""
    for idx in zones.get("acknowledgement", []):
        p = doc.paragraphs[idx]
        text = p.text.strip()
        if not text:
            continue

        if text == "致谢":
            for run in p.runs:
                set_run_font(run, cn_font=FONT_HEI, en_font=FONT_TIMES,
                            size=SIZE_SANHAO)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                line_spacing=BODY_LINE_SPACING)
        else:
            for run in p.runs:
                set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TIMES,
                            size=SIZE_XIAOSI)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                line_spacing=BODY_LINE_SPACING,
                                first_line_indent=BODY_INDENT)


def _run_replace_text(p, new_text):
    """替换段落的全部文本（保留 drawing 等非文本元素）。"""
    if not p.runs:
        return
    for run in p.runs:
        if run.text.strip():
            t_elems = run._element.findall(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            if t_elems:
                t_elems[0].text = new_text
                for te in t_elems[1:]:
                    te.text = ''
            else:
                run.text = new_text
            break
    found_first = False
    for run in p.runs:
        if run.text.strip() and not found_first:
            found_first = True
            continue
        if run.text.strip():
            t_elems = run._element.findall(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            for te in t_elems:
                te.text = ''


def _replace_text_in_paragraph(p, old_text, new_text):
    """在段落中替换指定文本（跨 run 替换）。"""
    full = ''.join(r.text for r in p.runs)
    if old_text not in full:
        return False
    new_full = full.replace(old_text, new_text)
    if p.runs:
        # 简单方案：所有文本写入第一个 run
        t_elems = p.runs[0]._element.findall(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        if t_elems:
            t_elems[0].text = new_full
        else:
            p.runs[0].text = new_full
    return True


def fix_figure_table_numbering(doc, zones):
    """修正图序表序、文中引用、公式格式。

    1. 图序: 图X.X.X 或 图X.X.X-Y → 图X-N
    2. 表序: 表X.X.X 或 表X.X.X-Y → 表X-N
    3. 文中引用: 如 "表X.X" → "表X-N" 同步修正
    4. 公式: 居中, 式序保持 (式X.Y) 或转为 (式X-N)
    """
    body_indices = zones.get("body", [])
    if not body_indices:
        return

    paragraphs = doc.paragraphs

    # Step 1: 找出章标题
    chapter_starts = []
    cn_num_map = {'一':1,'二':2,'三':3,'四':4,'五':5,'六':6,'七':7,'八':8,'九':9,'十':10}
    for idx in body_indices:
        p = paragraphs[idx]
        text = p.text.strip()
        if not _is_chapter_style(p.style.name, text):
            continue
        if text in ("参考文献", "致谢"):
            continue

        ch_num = 0
        # "第一章 xxx" or "第1章 xxx"
        m1 = re.match(r'^第\s*([一二三四五六七八九十\d]+)\s*章\b', text)
        if m1:
            ns = m1.group(1)
            ch_num = cn_num_map.get(ns, int(ns) if ns.isdigit() else 0)
        else:
            # "1 前言" / "1 xxx"
            m2 = re.match(r'^(\d+)\s+', text)
            if m2:
                ch_num = int(m2.group(1))

        if ch_num > 0:
            chapter_starts.append((idx, text, ch_num))

    if not chapter_starts:
        return

    # Step 2: 建立章节 → 图/表序号映射
    # 同时还建立旧编号 → 新编号的映射表
    old_to_new = {}   # "图2.4.2-1" → " 图2-?  XXX"

    for ci, (ch_start, ch_text, ch_num) in enumerate(chapter_starts):
        if ci + 1 < len(chapter_starts):
            ch_end = chapter_starts[ci + 1][0]
        else:
            ch_end = body_indices[-1] + 1

        fig_count = 0
        tbl_count = 0
        # 匹配图/表题注: 图/表 + 编号（可含.和-） + 可选空格 + 标题文字
        cap_pattern = re.compile(r'^(图|表)\s*([\d]+(?:[.\-][\d]+)*)\s*(.+)')

        for pi in range(ch_start + 1, ch_end):
            p = paragraphs[pi]
            text = p.text.strip()
            m = cap_pattern.match(text)
            if not m:
                continue

            prefix = m.group(1)
            old_num = m.group(2)
            rest = m.group(3)

            # 过滤误匹配：正文"表2.4判别结果显示..." → 是正文不是题注
            # 题注特征：以名词/专有名词/英文/符号开头，不是以动词/连词开头
            if re.match(r'^(判别|结果|显示|表明|可知|可以|得出|为不|可|本|该|其|与|并|且|但|而)', rest):
                continue
            # 题注通常30字以内
            if len(rest) > 30:
                continue

            if prefix == "图":
                fig_count += 1
                new_num = f"{ch_num}-{fig_count}"
            else:
                tbl_count += 1
                new_num = f"{ch_num}-{tbl_count}"

            old_key = f"{prefix}{old_num}"
            new_key = f"{prefix}{new_num}"
            old_to_new[old_key] = (new_key, rest)

            new_text = f"{new_key} {rest}"
            _run_replace_text(p, new_text)

    # Step 3: 修正文中引用
    if old_to_new:
        # 也加入带空格的变体（"表 4.1" → "表4-1"）
        expanded = {}
        for ok, (nk, _) in old_to_new.items():
            expanded[ok] = nk
            if re.match(r'^(图|表)(\d)', ok):
                m = re.match(r'^([图表])([\d].*)', ok)
                if m:
                    expanded[f"{m.group(1)} {m.group(2)}"] = nk

        # 添加短格式替代：对于 "图2.4.2-1" → 可能被简写为 "图2.4"
        short_forms = {}
        for ok, nk in list(expanded.items()):
            m = re.match(r'^([图表])([\d])-([\d]+)$', nk)
            if m:
                # 新编号 "图2-3" → 可能残留旧格式 "图2.X" 或 "图2.X.X"
                ch = m.group(2)
                pass
            # "图2.4.2-1" → short "图2.4"
            m = re.match(r'^([图表])([\d]+)\.([\d]+)', ok)
            if m:
                sf = f"{m.group(1)}{m.group(2)}.{m.group(3)}"
                if sf not in expanded:
                    short_forms[sf] = nk
        expanded.update(short_forms)

        # 清理"混合残留"：如 "图2-12.6" → "图2-12"，"表2-3.6.4-1" → "表2-3"
        for idx in body_indices:
            p = paragraphs[idx]
            full = ''.join(r.text for r in p.runs)
            full, n = re.subn(r'([图表]\d+-\d+)\.\d+(\.\d+)?(-\d+)?', r'\1', full)
            full, n2 = re.subn(r'至\s*[图表\s]*[\d.\-]+(?=\s|$|。|，|；)', r'', full)
            if (n + n2) > 0 and p.runs:
                _run_replace_text(p, full)

        sorted_keys = sorted(expanded.keys(), key=len, reverse=True)

        for idx in body_indices:
            p = paragraphs[idx]
            text = p.text.strip()
            if not text:
                continue

            full = ''.join(r.text for r in p.runs)
            modified = False
            for old_key in sorted_keys:
                if old_key in full:
                    new_key = expanded[old_key]
                    full = full.replace(old_key, new_key)
                    modified = True
            if modified and p.runs:
                _run_replace_text(p, full)

    # Step 4: 公式格式化（居中，式序靠右）
    # 式序写法: (式4.1), (式4.1), (4.1) 等 → 公式段落居中
    formula_pattern = re.compile(r'^[\(（]式\s*[\d.\-]+[\)）]$')
    for idx in body_indices:
        p = paragraphs[idx]
        text = p.text.strip()
        if formula_pattern.match(text):
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                line_spacing=BODY_LINE_SPACING)
            for run in p.runs:
                set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TIMES, size=SIZE_XIAOSI)


SIZE_XIAOER = Pt(18)   # 小二

def insert_toc(doc):
    """在摘要和正文之间插入自动目录域。

    在关键词段落后、第一个章标题前插入 TOC 域代码。
    用户打开 Word 后右键 -> 更新域 即可生成目录。
    """
    paragraphs = doc.paragraphs

    # 查找插入位置：关键词段落之后、第一个章标题之前
    toc_insert_after = None
    first_chapter_idx = None

    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if text.startswith("关键词"):
            toc_insert_after = i
        if toc_insert_after is not None and i > toc_insert_after and _is_chapter_style(p.style.name, text):
            first_chapter_idx = i
            break

    if toc_insert_after is None or first_chapter_idx is None:
        return

    # 检查是否已存在 TOC
    for i in range(toc_insert_after + 1, first_chapter_idx):
        instrs = paragraphs[i]._element.findall(
            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instrText')
        for ins in instrs:
            if ins.text and 'TOC' in ins.text.upper():
                return  # TOC already exists

    # 在关键词段落后插入空行 + TOC 域
    body_elem = paragraphs[toc_insert_after]._element.getparent()

    # 插入一个空行
    empty_p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    empty_p.append(pPr)
    body_elem.insert(list(body_elem).index(paragraphs[toc_insert_after]._element) + 1, empty_p)

    # 插入目录标题（小二黑体居中）
    toc_title = OxmlElement('w:p')
    tPPr = OxmlElement('w:pPr')
    tjc = OxmlElement('w:jc')
    tjc.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'center')
    tPPr.append(tjc)
    toc_title.append(tPPr)
    tR = OxmlElement('w:r')
    tRPr = OxmlElement('w:rPr')
    tRF = OxmlElement('w:rFonts')
    tRF.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', FONT_HEI)
    tRF.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', FONT_TIMES)
    tRPr.append(tRF)
    tSz = OxmlElement('w:sz')
    tSz.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(int(SIZE_XIAOER / 6350)))
    tRPr.append(tSz)
    tR.append(tRPr)
    tT = OxmlElement('w:t')
    tT.set('{http://schemas.xmlsoap.org/XML/1998/namespace}space', 'preserve')
    tT.text = '目  录'
    tR.append(tT)
    toc_title.append(tR)
    body_elem.insert(list(body_elem).index(paragraphs[toc_insert_after]._element) + 2, toc_title)

    # 插入 TOC 域
    toc_p = OxmlElement('w:p')
    tocPPr = OxmlElement('w:pPr')
    toc_p.append(tocPPr)

    # begin fldChar
    r1 = OxmlElement('w:r')
    fld1 = OxmlElement('w:fldChar')
    fld1.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'begin')
    r1.append(fld1)
    toc_p.append(r1)

    # instrText
    r2 = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set('{http://schemas.xmlsoap.org/XML/1998/namespace}space', 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2.append(instr)
    toc_p.append(r2)

    # separate
    r3 = OxmlElement('w:r')
    sep = OxmlElement('w:fldChar')
    sep.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'separate')
    r3.append(sep)
    toc_p.append(r3)

    # placeholder text
    r4 = OxmlElement('w:r')
    pt = OxmlElement('w:t')
    pt.text = '（请在 Word 中右键此处 → 更新域，以生成目录）'
    r4.append(pt)
    toc_p.append(r4)

    # end fldChar
    r5 = OxmlElement('w:r')
    fld2 = OxmlElement('w:fldChar')
    fld2.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'end')
    r5.append(fld2)
    toc_p.append(r5)

    body_elem.insert(list(body_elem).index(paragraphs[toc_insert_after]._element) + 3, toc_p)

    # 插入空行
    empty_p2 = OxmlElement('w:p')
    body_elem.insert(list(body_elem).index(paragraphs[toc_insert_after]._element) + 4, empty_p2)

    # 在目录后、第一章标题前插入分页符
    ch_elem = paragraphs[first_chapter_idx]._element
    ch_pPr = ch_elem.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
    if ch_pPr is None:
        ch_pPr = OxmlElement('w:pPr')
        ch_elem.insert(0, ch_pPr)
    pb = ch_pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pageBreakBefore')
    if pb is None:
        pb = OxmlElement('w:pageBreakBefore')
        ch_pPr.append(pb)


def _apply_formula_tab_stops(p):
    """设置公式段落的制表位：居中制表位 + 右对齐制表位。

    参考袁颖论文的公式样式：center tab at ~20 char, right tab at ~39 char。
    """
    pPr = p._element.get_or_add_pPr()
    # 先移除已有 tabs
    old_tabs = pPr.findall(qn('w:tabs'))
    for t in old_tabs:
        pPr.remove(t)

    tabs = OxmlElement('w:tabs')
    # 居中制表位
    ct = OxmlElement('w:tab')
    ct.set(qn('w:val'), 'center')
    ct.set(qn('w:pos'), '4536')  # ~315pt, page center for 3.2cm margins
    tabs.append(ct)
    # 右制表位（式序靠右）
    rt = OxmlElement('w:tab')
    rt.set(qn('w:val'), 'right')
    rt.set(qn('w:pos'), '9072')  # ~630pt, right margin
    tabs.append(rt)
    pPr.append(tabs)


def _cleanup_mixed_refs(doc, zones):
    """最后清理所有混合残留引用：图2-12.6 → 图2-12，表2-3.6.4-1 → 表2-3 等。"""
    body_indices = zones.get("body", [])
    paragraphs = doc.paragraphs
    for idx in body_indices:
        p = paragraphs[idx]
        full = ''.join(r.text for r in p.runs)
        if not full.strip():
            continue
        modified = False
        # "图2-12.6" / "表2-3.6" / "表2-3.6.4-1" → clean suffix
        full, n = re.subn(r'([图表]\d+-\d+)\.\d+(\.\d+)?(-\d+)?', r'\1', full)
        if n > 0: modified = True
        # "至2.4.6.2-11" → remove
        full, n2 = re.subn(r'至\s*[图表\s]*[\d.\-]+', r'', full)
        if n2 > 0: modified = True
        if modified and p.runs:
            _run_replace_text(p, full)


def _is_formula_content(p):
    """判断段落是否为公式内容（非标题、非正文段落）。

    公式内容特征：以数学符号/变量开头，或上一段是(式X.Y)标签段落。
    """
    text = p.text.strip()
    if not text:
        return False
    # 不是章/节/小节标题，不是图表题注，不是普通长文本正文
    if _is_chapter_style(p.style.name, text):
        return False
    if _is_section_style(p.style.name, text):
        return False
    if _is_subsection_style(p.style.name, text):
        return False
    if _is_caption_style(p.style.name):
        return False
    if re.match(r'^[\(（]式', text):
        return False
    # 文字太长的是正文
    if len(text) > 80:
        return False
    return True


def format_formulas(doc, zones):
    """格式化数学公式段落：合并公式内容 + 式序到同一段，制表位居中 + 式序靠右。

    原作者常把公式内容和式序分两段写。此函数将它们合并为正确格式：
    {tab}公式内容{tab}（式X.Y）
    """
    body_indices = zones.get("body", [])
    paragraphs = doc.paragraphs
    eq_label_pat = re.compile(r'^[\(（]式\s*[\d.\-]+[\)）]$')

    # Pass 1: 找到式序号段落，合并上一段
    merge_pairs = []
    for idx in body_indices:
        p = paragraphs[idx]
        text = p.text.strip()
        if not eq_label_pat.match(text):
            continue
        if idx == 0:
            continue
        prev = paragraphs[idx - 1]
        prev_text = prev.text.strip()
        # 跳过：上一段也是式序、空段、长正文(>80字)
        if not prev_text:
            continue
        if eq_label_pat.match(prev_text):
            continue
        if len(prev_text) > 80:
            continue
        # 上一段是图表题注也跳过
        if re.match(r'^(图|表)\d', prev_text):
            continue
        merge_pairs.append((idx - 1, idx, text))

    # 从后往前合并
    for prev_idx, label_idx, label_text in reversed(merge_pairs):
        prev = paragraphs[prev_idx]
        label_p = paragraphs[label_idx]

        # 获取公式内容文本
        content_text = ''.join(r.text for r in prev.runs).strip()

        # 用 OxmlElement 创建新的干净 run
        # 完全清除 label_p 中的所有子元素（包括旧 run 和 pPr）
        for child in list(label_p._element):
            if child.tag != '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr':
                label_p._element.remove(child)

        # 创建新 run: tab + content + tab + label
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), FONT_SONG)
        rFonts.set(qn('w:ascii'), FONT_TIMES)
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(SIZE_XIAOSI / 6350)))
        rPr.append(sz)
        new_run.append(rPr)

        # tab 元素
        tab1 = OxmlElement('w:tab')
        new_run.append(tab1)
        # 公式内容
        t1 = OxmlElement('w:t')
        t1.set('{http://schemas.xmlsoap.org/XML/1998/namespace}space', 'preserve')
        t1.text = ' ' + content_text + ' '
        new_run.append(t1)
        # tab
        tab2 = OxmlElement('w:tab')
        new_run.append(tab2)
        # 式序
        t2 = OxmlElement('w:t')
        t2.set('{http://schemas.xmlsoap.org/XML/1998/namespace}space', 'preserve')
        t2.text = label_text
        new_run.append(t2)

        label_p._element.append(new_run)

        # 删除公式内容段
        prev._element.getparent().remove(prev._element)

        # 重建 paragraphs 引用
        paragraphs = doc.paragraphs

    # Pass 2: 格式化所有公式段落
    paragraphs = doc.paragraphs
    body_indices = zones.get("body", [])
    for idx in body_indices:
        if idx >= len(paragraphs):
            continue
        p = paragraphs[idx]
        text = p.text.strip()
        if eq_label_pat.match(text):
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                line_spacing=BODY_LINE_SPACING,
                                first_line_indent=None)
            _apply_formula_tab_stops(p)
            for run in p.runs:
                set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TIMES,
                            size=SIZE_XIAOSI)


def format_tables(doc):
    """修正正文表格内文字格式（跳过封面表格）。"""
    for ti, table in enumerate(doc.tables):
        if ti == 0:
            continue  # 封面表格已在 format_cover 中处理
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if not text:
                        continue
                    for run in paragraph.runs:
                        set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TIMES,
                                    size=SIZE_XIAOSI)
                    set_paragraph_format(paragraph, line_spacing=BODY_LINE_SPACING)


def format_page_numbers(doc):
    """添加页码（页脚）。"""
    sections = doc.sections
    for i, section in enumerate(sections):
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()

        # 清除现有页脚内容
        for p in footer.paragraphs:
            for run in p.runs:
                run.text = ""

        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TIMES,
                    size=SIZE_XIAOWU)

        # 插入页码域代码
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)

        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)

        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)


def format_headers_footers(doc):
    """修正页眉页脚字体。"""
    for section in doc.sections:
        # 页眉
        if section.header:
            for p in section.header.paragraphs:
                if p.text.strip():
                    for run in p.runs:
                        set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TIMES,
                                    size=SIZE_XIAOWU)
                    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # 页脚
        if section.footer:
            for p in section.footer.paragraphs:
                if p.text.strip():
                    for run in p.runs:
                        set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TIMES,
                                    size=SIZE_XIAOWU)


def main():
    if len(sys.argv) < 3:
        print("用法: uv run --with python-docx python3 scripts/formatter.py input.docx output.docx")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    if not Path(input_path).exists():
        print(f"错误: 找不到输入文件 {input_path}")
        sys.exit(1)

    print(f"正在读取: {input_path}")
    doc = Document(input_path)

    print("正在修正格式...")

    # 1. 页面设置
    print("  [1/10] 页面设置...")
    format_page_setup(doc)

    # 2. 分析文档结构
    print("  [2/10] 分析文档结构...")
    zones = get_structure_zones(doc)

    # 3. 封面格式
    print("  [3/10] 封面格式...")
    format_cover(doc, zones)

    # 4. 图序表序修正 + 公式格式
    print("  [4/10] 图序表序修正...")
    fix_figure_table_numbering(doc, zones)

    # 5. 摘要格式
    print("  [5/10] 摘要格式...")
    format_abstract(doc, zones)

    # 6. 标题前空行 + 正文格式（含章标题分页）
    print("  [6/10] 标题前空行...")
    _add_blank_before_headings(doc, zones)
    # 重建 zones（段落被插入后索引已变）
    zones = get_structure_zones(doc)
    print("  [7/10] 正文格式...")
    format_body(doc, zones)

    # 8. 致谢格式
    print("  [8/10] 致谢格式...")
    format_acknowledgement(doc, zones)

    # 9. 公式格式修正（在正文格式之后，避免被覆盖）
    print("  [9/10] 公式格式...")
    format_formulas(doc, zones)

    # 9b. 最终清理混合残留引用
    print("  [9b/10] 清理残留引用...")
    _cleanup_mixed_refs(doc, zones)

    # 10. 插入目录域
    print("  [10/10] 插入目录域...")
    insert_toc(doc)

    # 11. 表格格式 + 页眉页脚
    print("  [11/11] 表格格式...")
    format_tables(doc)
    format_headers_footers(doc)

    # 12. 页码
    print("  [12/12] 页码...")
    format_page_numbers(doc)

    # 保存
    doc.save(output_path)
    print(f"完成! 已保存到: {output_path}")


if __name__ == "__main__":
    main()
