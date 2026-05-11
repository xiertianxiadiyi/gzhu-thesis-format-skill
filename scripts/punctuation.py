"""
标点符号修复工具（适配广州大学毕业论文）

修复中英文标点混用问题，适用于学术论文场景。
"""

import re
import sys
from pathlib import Path

from docx import Document


def has_chinese(text):
    """检查文本是否包含中文字符。"""
    return bool(re.search(r'[\u4e00-\u9fff\u3400-\u4dbf]', text))


def fix_punctuation_in_text(text):
    """修复段落文本中的标点符号。"""
    if not text:
        return text

    # 只在包含中文的段落中处理
    if not has_chinese(text):
        return text

    # 1. 省略号
    text = re.sub(r'(?<![。．\.])\.\.(?![。．\.])', '……', text)
    text = re.sub(r'(?<![。．])。。', '……', text)

    # 2. 破折号
    text = re.sub(r'--+', '——', text)
    text = re.sub(r'—(?!—)', '——', text)

    # 3. 括号
    text = re.sub(r'(?<=[\u4e00-\u9fff])\(', '（', text)
    text = re.sub(r'\)(?=[\u4e00-\u9fff》」）])', '）', text)

    # 4. 冒号（中文上下文）
    text = re.sub(r'(?<=[\u4e00-\u9fff]):', '：', text)
    text = re.sub(r':(?=[\u4e00-\u9fff\s])', '：', text)

    # 5. 分号（中文上下文）
    text = re.sub(r'(?<=[\u4e00-\u9fff]);', '；', text)
    text = re.sub(r';(?=[\u4e00-\u9fff])', '；', text)

    # 6. 问号（中文上下文）
    text = re.sub(r'(?<=[\u4e00-\u9fff])\?', '？', text)

    # 7. 感叹号（中文上下文）
    text = re.sub(r'(?<=[\u4e00-\u9fff])!', '！', text)

    # 8. 逗号（中文上下文）
    text = re.sub(r'(?<=[\u4e00-\u9fff]),', '，', text)
    text = re.sub(r',(?=[\u4e00-\u9fff])', '，', text)

    # 9. 句号（中文上下文）
    text = re.sub(r'(?<=[\u4e00-\u9fff])\.', '。', text)
    text = re.sub(r'\.(?=[\u4e00-\u9fff])', '。', text)

    # 10. 中文引号配对（将直引号转为弯引号）
    # 先处理引号内无嵌套的情况
    parts = []
    i = 0
    in_quote = False
    while i < len(text):
        if text[i] == '"':
            if in_quote:
                parts.append('\u201d')
                in_quote = False
            else:
                parts.append('\u201c')
                in_quote = True
        elif text[i] == "'" and has_chinese(text[max(0, i-1):min(len(text), i+2)]):
            if i > 0 and text[i-1] in '（(':
                parts.append('\u2018')
                in_quote = True
            elif i < len(text) - 1 and text[i+1] in '）)' :
                parts.append('\u2019')
                in_quote = False
            else:
                parts.append(text[i])
        else:
            parts.append(text[i])
        i += 1

    return ''.join(parts)


def fix_paragraph(p):
    """修复单个段落的标点。"""
    if not p.text.strip():
        return

    # 合并所有 run 的文本
    full_text = ''.join(run.text for run in p.runs)
    if not full_text.strip():
        return

    fixed_text = fix_punctuation_in_text(full_text)

    # 将修复后的文本写回第一个 run，清空其他 run
    if p.runs:
        p.runs[0].text = fixed_text
        for run in p.runs[1:]:
            run.text = ''


def fix_table_cells(doc):
    """修复表格单元格内的标点。"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    fix_paragraph(p)


def main():
    if len(sys.argv) < 3:
        print("用法: uv run --with python-docx python3 scripts/punctuation.py input.docx output.docx")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    if not Path(input_path).exists():
        print(f"错误: 找不到输入文件 {input_path}")
        sys.exit(1)

    doc = Document(input_path)

    print("正在修复标点符号...")

    count = 0
    for p in doc.paragraphs:
        old_text = p.text
        fix_paragraph(p)
        if p.text != old_text and p.text.strip():
            count += 1

    fix_table_cells(doc)

    doc.save(output_path)
    print(f"完成! 修复了 {count} 个段落，已保存到: {output_path}")


if __name__ == "__main__":
    main()
